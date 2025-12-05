import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from datetime import datetime
from PIL import Image
import pytesseract
import fitz  
from docx import Document
from pdf2image import convert_from_path
import numpy as np
from openai import OpenAI
import threading
import webbrowser


OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if OPENAI_API_KEY:
    client = OpenAI(api_key=OPENAI_API_KEY)
else:
    client = None


def simple_sent_tokenize(text):
    import re
    sentences = re.split(r'(?<=[.!?\n])\s+', text)
    return [s.strip() for s in sentences if s.strip()]

def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in [".txt", ".log"]:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif ext == ".docx":
            doc = Document(path)
            text = "\n".join([p.text for p in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " ".join(cell.text for cell in row.cells)
            return text
        elif ext == ".pdf":
            text = ""
            if fitz:
                with fitz.open(path) as pdf:
                    for page in pdf:
                        text += page.get_text("text") or ""
            if not text.strip() and convert_from_path and pytesseract:
                pages = convert_from_path(path)
                for page in pages:
                    text += pytesseract.image_to_string(page)
            return text
        elif ext in [".png", ".jpg", ".jpeg", ".bmp"] and pytesseract:
            img = Image.open(path)
            return pytesseract.image_to_string(img)
        else:
            return ""
    except:
        return ""

def get_embedding(text):
    if not client:
        return np.zeros(1536)
    try:
        resp = client.embeddings.create(
            model="text-embedding-3-large",
            input=text
        )
        return np.array(resp.data[0].embedding)
    except Exception as e:
        print("Embedding Error:", e)
        return np.zeros(1536)

def index_folder(folder, progress_callback=None):
    index = []
    total_files = sum(len(files) for _, _, files in os.walk(folder))
    count = 0
    for root, _, files in os.walk(folder):
        for file in files:
            count += 1
            path = os.path.join(root, file)
            text = extract_text(path)
            if not text.strip():
                continue
            for sentence in simple_sent_tokenize(text):
                if sentence:
                    emb = get_embedding(sentence)
                    index.append((path, sentence, emb))
            if progress_callback:
                progress_callback(f"[{count}/{total_files}] Indexed: {file}")
    return index

def semantic_search(index, query, top_k=10):
    if client is None:
        return []
    query_emb = get_embedding(query)
    results = []
    for path, sentence, emb in index:
        score = np.dot(query_emb, emb) / (np.linalg.norm(query_emb) * np.linalg.norm(emb))
        results.append((score, path, sentence))
    results.sort(reverse=True)
    return results[:top_k]


class ContentFilterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AI-Powered Content Filtering App")
        self.root.geometry("1000x750")
        self.create_ui()

    def create_ui(self):
      
        bg_color = "#6C2BD0"       
        frame_bg = "#ffffff"        
        accent_color = "#33C3F0" 
        button_color = "#33C3F0"    
        entry_bg = "#f0f0f0"       
        report_bg = "#fefefe"       
        font_header = ("Segoe UI", 11, "bold")
        font_normal = ("Segoe UI", 10)

        self.root.configure(bg=bg_color)

    
        tk.Label(self.root, text="FOLDER DIRECTORY", font=font_header, bg=bg_color, fg="#ffffff").pack(anchor="w", padx=20, pady=(10,2))
        top_frame = tk.Frame(self.root, bg=bg_color)
        top_frame.pack(fill="x", padx=20)
        self.folder_entry = tk.Entry(top_frame, font=font_normal, width=50, bg=entry_bg, relief="flat")
        self.folder_entry.pack(side="left", padx=5, pady=5, ipady=3)
        tk.Button(top_frame, text="Browse", command=self.select_folder, bg=accent_color, fg="white", relief="flat").pack(side="left", padx=5, pady=5)

        
        tk.Label(self.root, text="SEARCH CONTEXT", font=font_header, bg=bg_color, fg="#ffffff").pack(anchor="w", padx=20, pady=(15,3))
        self.search_entry = tk.Entry(self.root, font=font_normal, width=50, bg=entry_bg, relief="flat")
        self.search_entry.pack(padx=20, fill="x")
        tk.Button(self.root, text="Search", command=self.start_search_thread, bg=button_color, fg="white", relief="flat", width=12).pack(anchor="w", padx=20, pady=10)

        
        tk.Label(self.root, text="RESULT TABLE", font=font_header, bg=bg_color, fg="#ffffff").pack(anchor="w", padx=20)
        table_frame = tk.Frame(self.root, bg=frame_bg)
        table_frame.pack(padx=20, pady=5, fill="both", expand=True)
        columns = ("score", "path", "line")
        self.table = ttk.Treeview(table_frame, columns=columns, show="headings", height=12)
        style = ttk.Style()
        style.theme_use("default")
        style.configure("Treeview", background=frame_bg, foreground="#333333", rowheight=25, fieldbackground=frame_bg, font=font_normal)
        style.configure("Treeview.Heading", font=font_header, background=accent_color, foreground="white")
        self.table.heading("score", text="Score")
        self.table.heading("path", text="File Path")
        self.table.heading("line", text="Matched Line")
        self.table.column("score", width=80)
        self.table.column("path", width=500)
        self.table.column("line", width=400)
        self.table.pack(fill="both", expand=True)
        self.table.bind("<Double-1>", self.open_file)

        
        tk.Label(self.root, text="REPORT", font=font_header, bg=bg_color, fg="#ffffff").pack(anchor="w", padx=20, pady=(10,0))
        self.report_box = scrolledtext.ScrolledText(self.root, height=12, font=font_normal, bg=report_bg, relief="flat")
        self.report_box.pack(padx=20, pady=10, fill="both", expand=False)

    
    def select_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.folder_entry.delete(0, tk.END)
            self.folder_entry.insert(0, folder)


    def start_search_thread(self):
        thread = threading.Thread(target=self.perform_search)
        thread.start()

    def perform_search(self):
        folder = self.folder_entry.get().strip()
        query = self.search_entry.get().strip()
        if not folder or not os.path.isdir(folder):
            messagebox.showerror("Error", "Please select a valid folder.")
            return
        if not query:
            messagebox.showwarning("Warning", "Enter a search term.")
            return

        self.table.delete(*self.table.get_children())
        self.report_box.delete("1.0", tk.END)

        def update_progress(msg):
            self.report_box.insert(tk.END, msg + "\n")
            self.report_box.see(tk.END)
            self.root.update()

        update_progress("Indexing and performing search... Please wait.")
        index = index_folder(folder, progress_callback=update_progress)
        
        if client:
            results = semantic_search(index, query, top_k=20)
        else:
            results = []
            for path, sentence, _ in index:
                if query.lower() in sentence.lower():
                    results.append((1.0, path, sentence))

        if results:
            for score, path, line in results:
                self.table.insert("", tk.END, values=(f"{score:.3f}", path, line))
            update_progress(f"{len(results)} results found.")
        else:
            update_progress("No matches found.")

    
    def open_file(self, event):
        selected = self.table.focus()
        if not selected:
            return
        file_path = self.table.item(selected)["values"][1]
        if os.path.exists(file_path):
            os.startfile(file_path)
        else:
            messagebox.showerror("Error", "File not found.")

if __name__ == "__main__":
    root = tk.Tk()
    app = ContentFilterApp(root)
    root.mainloop()
